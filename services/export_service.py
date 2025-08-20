"""
Export Service - Handles file exports for tabular data and reports
"""

import os
import pandas as pd
import json
from datetime import datetime
from typing import Dict, List, Any, Optional
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import xlsxwriter

class ExportService:
    def __init__(self):
        self.export_dir = "exports"
        self._ensure_export_directory()
    
    def _ensure_export_directory(self):
        """Create export directory if it doesn't exist"""
        if not os.path.exists(self.export_dir):
            os.makedirs(self.export_dir)
    
    def detect_tabular_query(self, query: str, response: str) -> bool:
        """Detect if the query/response should be exported as a table"""
        tabular_keywords = [
            'compare', 'comparison', 'vs', 'versus', 'table', 'list',
            'ranking', 'rank', 'top', 'best', 'worst', 'summary',
            'overview', 'statistics', 'stats', 'data', 'numbers',
            'breakdown', 'analysis', 'report', 'excel', 'download',
            'downloadable', 'export', 'file', 'format', 'csv',
            'word', 'doc', 'all comments', 'all 500', 'comments on',
            'give me all', 'show me all', 'list all', 'display all'
        ]
        
        query_lower = query.lower()
        response_lower = response.lower()
        
        # Check for tabular keywords
        keyword_found = any(keyword in query_lower for keyword in tabular_keywords)
        
        # Check for specific export indicators
        export_indicators = ['excel', 'download', 'export', 'file', 'format']
        export_requested = any(indicator in query_lower for indicator in export_indicators)
        
        # Check for table-like structure in response
        has_table_structure = ('|' in response and '-' in response) or \
                            (response.count('\n') > 5 and ':' in response)
        
        # Check for specific comment requests
        comment_request = any(phrase in query_lower for phrase in [
            'all comments', 'all 500', 'comments on', 'comments for',
            'show comments', 'list comments', 'give me comments'
        ])
        
        return keyword_found or has_table_structure or export_requested or comment_request
    
    def create_excel_export(self, data: Dict[str, Any], filename: str = None) -> str:
        """Create Excel file from data"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"solysai_export_{timestamp}.xlsx"
        
        filepath = os.path.join(self.export_dir, filename)
        
        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            # Main analysis sheet
            if 'analysis' in data:
                analysis_df = self._format_analysis_for_excel(data['analysis'])
                analysis_df.to_excel(writer, sheet_name='Analysis', index=False)
                self._style_excel_sheet(writer.sheets['Analysis'])
            
            # Comments data sheet
            if 'comments_data' in data:
                comments_df = self._format_comments_for_excel(data['comments_data'])
                comments_df.to_excel(writer, sheet_name='Comments_Data', index=False)
                self._style_excel_sheet(writer.sheets['Comments_Data'])
            
            # Summary statistics
            if 'statistics' in data:
                stats_df = pd.DataFrame(data['statistics'])
                stats_df.to_excel(writer, sheet_name='Statistics', index=False)
                self._style_excel_sheet(writer.sheets['Statistics'])
            
            # Sources sheet
            if 'sources' in data:
                sources_df = pd.DataFrame(data['sources'])
                sources_df.to_excel(writer, sheet_name='Sources', index=False)
                self._style_excel_sheet(writer.sheets['Sources'])
        
        return filepath
    
    def create_word_export(self, data: Dict[str, Any], filename: str = None) -> str:
        """Create Word document from data"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"solysai_report_{timestamp}.docx"
        
        filepath = os.path.join(self.export_dir, filename)
        
        doc = Document()
        
        # Title
        title = doc.add_heading('SolysAI Market Intelligence Report', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Metadata
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph(f"Query: {data.get('query', 'N/A')}")
        doc.add_page_break()
        
        # Executive Summary
        if 'summary' in data:
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(data['summary'])
        
        # Main Analysis
        if 'analysis' in data:
            doc.add_heading('Detailed Analysis', level=1)
            doc.add_paragraph(data['analysis'])
        
        # Key Findings
        if 'key_findings' in data:
            doc.add_heading('Key Findings', level=1)
            for finding in data['key_findings']:
                p = doc.add_paragraph(finding, style='List Bullet')
        
        # Data Tables
        if 'tables' in data:
            doc.add_heading('Data Tables', level=1)
            for table_name, table_data in data['tables'].items():
                doc.add_heading(table_name, level=2)
                if isinstance(table_data, pd.DataFrame):
                    self._add_dataframe_to_doc(doc, table_data)
                else:
                    doc.add_paragraph(str(table_data))
        
        # Sources
        if 'sources' in data:
            doc.add_page_break()
            doc.add_heading('Sources', level=1)
            for i, source in enumerate(data['sources'], 1):
                doc.add_paragraph(f"{i}. {source.get('title', 'Unknown')}")
                doc.add_paragraph(f"   URL: {source.get('url', 'N/A')}")
                if 'snippet' in source:
                    doc.add_paragraph(f"   Summary: {source['snippet'][:200]}...")
        
        doc.save(filepath)
        return filepath
    
    def _format_analysis_for_excel(self, analysis_text: str) -> pd.DataFrame:
        """Convert analysis text to structured DataFrame"""
        lines = analysis_text.split('\n')
        data = []
        
        current_section = "General"
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('#') or line.startswith('**'):
                current_section = line.replace('#', '').replace('**', '').strip()
            else:
                data.append({
                    'Section': current_section,
                    'Content': line,
                    'Type': 'Analysis'
                })
        
        return pd.DataFrame(data)
    
    def _format_comments_for_excel(self, comments_data: List[Dict]) -> pd.DataFrame:
        """Format comments data for Excel export with ADVANCED sentiment classification"""
        formatted_data = []
        
        for comment in comments_data:
            # Use ADVANCED sentiment classification if available
            sentiment_info = comment.get('sentiment_classification', {})
            advanced_sentiment_info = comment.get('advanced_sentiment_classification', {})
            
            # Prioritize advanced classification
            if advanced_sentiment_info:
                sentiment = advanced_sentiment_info.get('sentiment', 'Neutral').title()
                confidence = advanced_sentiment_info.get('confidence', 0.0)
                sarcasm_detected = advanced_sentiment_info.get('sarcasm_detected', False)
                language_mix = advanced_sentiment_info.get('language_analysis', {}).get('is_mixed', False)
                context = advanced_sentiment_info.get('context', 'general')
                product_relevance = advanced_sentiment_info.get('product_relevance', 'unknown')
                company_mentions = advanced_sentiment_info.get('company_analysis', {}).get('primary_company', 'None')
                
                # Enhanced categorization with context
                category = self._get_advanced_category(context, advanced_sentiment_info)
                
                formatted_comment = {
                    'OEM': comment.get('oem', 'Unknown'),
                    'Comment': comment.get('text', ''),
                    'Author': comment.get('author', 'Anonymous'),
                    'Likes': comment.get('likes', 0),
                    'Date': comment.get('date', 'Unknown'),
                    'Video_Title': comment.get('video_title', ''),
                    'Video_URL': comment.get('video_url', ''),
                    'Sentiment': sentiment,
                    'Sentiment_Confidence': round(confidence, 3),
                    'Sarcasm_Detected': 'Yes' if sarcasm_detected else 'No',
                    'Language_Mixed': 'Yes' if language_mix else 'No',
                    'Category': category,
                    'Context': context.title(),
                    'Product_Relevance': product_relevance.title(),
                    'Company_Mentions': company_mentions or 'None',
                    'Analysis_Method': 'Advanced_Multi_Layer'
                }
                
                # Add emoji information if available
                emoji_info = advanced_sentiment_info.get('emoji_analysis', {})
                if emoji_info.get('has_emojis', False):
                    formatted_comment['Emojis_Found'] = emoji_info.get('emoji_count', 0)
                    formatted_comment['Emoji_Sentiment'] = emoji_info.get('emoji_sentiment', 'neutral').title()
                else:
                    formatted_comment['Emojis_Found'] = 0
                    formatted_comment['Emoji_Sentiment'] = 'None'
                
            elif sentiment_info:
                # Fallback to regular sentiment classification
                sentiment = sentiment_info.get('sentiment', 'Neutral').title()
                confidence = sentiment_info.get('confidence', 0.0)
                context = sentiment_info.get('context', 'general')
                
                formatted_comment = {
                    'OEM': comment.get('oem', 'Unknown'),
                    'Comment': comment.get('text', ''),
                    'Author': comment.get('author', 'Anonymous'),
                    'Likes': comment.get('likes', 0),
                    'Date': comment.get('date', 'Unknown'),
                    'Video_Title': comment.get('video_title', ''),
                    'Video_URL': comment.get('video_url', ''),
                    'Sentiment': sentiment,
                    'Sentiment_Confidence': round(confidence, 3),
                    'Category': self._categorize_comment_enhanced(comment.get('text', ''), context),
                    'Analysis_Method': 'Enhanced_Rules'
                }
            else:
                # Legacy fallback for old data
                formatted_comment = {
                    'OEM': comment.get('oem', 'Unknown'),
                    'Comment': comment.get('text', ''),
                    'Author': comment.get('author', 'Anonymous'),
                    'Likes': comment.get('likes', 0),
                    'Date': comment.get('date', 'Unknown'),
                    'Video_Title': comment.get('video_title', ''),
                    'Video_URL': comment.get('video_url', ''),
                    'Sentiment': self._analyze_sentiment_basic(comment.get('text', '')),
                    'Category': self._categorize_comment(comment.get('text', '')),
                    'Analysis_Method': 'Basic_Keywords'
                }
            
            formatted_data.append(formatted_comment)
        
        return pd.DataFrame(formatted_data)
    
    def _style_excel_sheet(self, worksheet):
        """Apply professional styling to Excel worksheet"""
        # Header styling
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        
        # Apply header styling
        for cell in worksheet[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
        
        # Auto-adjust column widths
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            
            adjusted_width = min(max_length + 2, 50)
            worksheet.column_dimensions[column_letter].width = adjusted_width
        
        # Add borders
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        for row in worksheet.iter_rows():
            for cell in row:
                cell.border = thin_border
    
    def _add_dataframe_to_doc(self, doc, df: pd.DataFrame):
        """Add DataFrame as table to Word document"""
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, column in enumerate(df.columns):
            header_cells[i].text = str(column)
        
        # Add data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
    
    def _analyze_sentiment_basic(self, text: str) -> str:
        """Simple sentiment analysis (legacy fallback)"""
        positive_words = ['good', 'great', 'excellent', 'amazing', 'love', 'best', 'fantastic']
        negative_words = ['bad', 'terrible', 'worst', 'hate', 'problem', 'issue', 'poor']
        
        text_lower = text.lower()
        positive_count = sum(1 for word in positive_words if word in text_lower)
        negative_count = sum(1 for word in negative_words if word in text_lower)
        
        if positive_count > negative_count:
            return 'Positive'
        elif negative_count > positive_count:
            return 'Negative'
        else:
            return 'Neutral'
    
    def _get_advanced_category(self, context: str, advanced_info: Dict) -> str:
        """Get enhanced category based on advanced classification"""
        # Context-based categorization
        context_mapping = {
            'service': 'Service & Support',
            'battery_performance': 'Battery & Range',
            'riding_experience': 'Performance & Experience', 
            'purchase_decision': 'Purchase & Pricing',
            'comparison': 'Product Comparison',
            'build_quality': 'Build Quality & Design',
            'features': 'Features & Technology',
            'financial': 'Pricing & Finance',
            'general': 'General Discussion'
        }
        
        category = context_mapping.get(context, 'General Discussion')
        
        # Add special modifiers based on advanced analysis
        if advanced_info.get('sarcasm_detected', False):
            category += ' (Sarcastic)'
        
        language_info = advanced_info.get('language_analysis', {})
        if language_info.get('is_mixed', False):
            category += ' (Multilingual)'
        
        return category
    
    def _categorize_comment_enhanced(self, text: str, context: str = 'general') -> str:
        """Enhanced categorization with context information"""
        if context == 'service':
            return 'Service & Support'
        elif context == 'product':
            return 'Product Review'
        elif context == 'experience':
            return 'User Experience'
        elif context == 'financial':
            return 'Pricing & Finance'
        else:
            return self._categorize_comment(text)
    
    def _categorize_comment(self, text: str) -> str:
        """Categorize comment by topic"""
        categories = {
            'Service': ['service', 'support', 'maintenance', 'repair'],
            'Battery': ['battery', 'range', 'charging', 'charge'],
            'Performance': ['speed', 'acceleration', 'performance', 'power'],
            'Price': ['price', 'cost', 'expensive', 'value', 'money'],
            'Quality': ['quality', 'build', 'reliability', 'durable'],
            'Features': ['features', 'technology', 'app', 'connectivity']
        }
        
        text_lower = text.lower()
        for category, keywords in categories.items():
            if any(keyword in text_lower for keyword in keywords):
                return category
        
        return 'General'
    
    def extract_table_from_response(self, response: str) -> Optional[pd.DataFrame]:
        """Extract table data from response text"""
        lines = response.split('\n')
        table_lines = []
        
        for line in lines:
            if '|' in line and len(line.split('|')) > 2:
                table_lines.append(line)
        
        if len(table_lines) < 2:
            return None
        
        # Parse table
        headers = [col.strip() for col in table_lines[0].split('|')[1:-1]]
        rows = []
        
        for line in table_lines[2:]:  # Skip header separator line
            if '|' in line:
                row = [col.strip() for col in line.split('|')[1:-1]]
                if len(row) == len(headers):
                    rows.append(row)
        
        if rows:
            return pd.DataFrame(rows, columns=headers)
        
        return None
    
    def get_export_stats(self) -> Dict[str, Any]:
        """Get statistics about exported files"""
        export_files = [f for f in os.listdir(self.export_dir) if f.endswith(('.xlsx', '.docx'))]
        
        return {
            'total_exports': len(export_files),
            'excel_files': len([f for f in export_files if f.endswith('.xlsx')]),
            'word_files': len([f for f in export_files if f.endswith('.docx')]),
            'latest_export': max(export_files, key=lambda x: os.path.getctime(os.path.join(self.export_dir, x))) if export_files else None
        }
